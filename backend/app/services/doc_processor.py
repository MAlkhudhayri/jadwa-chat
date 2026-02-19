"""
Document processing service using Docling.

Converts PDF, DOCX, PPTX into structured chunks with rich metadata.
Tables with numerical data are extracted separately for SQLite storage.
"""

import asyncio
import logging
import re
import tempfile
import os
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

logger = logging.getLogger(__name__)


@dataclass
class DocChunk:
    """A chunk of text from a document, with metadata."""
    text: str
    source: str                    # filename
    page: Optional[int] = None
    section: Optional[str] = None
    chunk_type: str = "text"       # text | table | heading
    chunk_index: int = 0
    total_chunks: int = 0


@dataclass
class ExtractedTable:
    """A table extracted from a document, potentially for SQLite."""
    headers: List[str]
    rows: List[List[str]]
    caption: Optional[str] = None
    page: Optional[int] = None
    source: str = ""


@dataclass
class DoclingResult:
    """Result of processing a document with Docling."""
    chunks: List[DocChunk] = field(default_factory=list)
    tables: List[ExtractedTable] = field(default_factory=list)
    title: Optional[str] = None
    num_pages: int = 0


class DoclingProcessor:
    """Process documents using Docling for rich structure extraction."""

    def __init__(self):
        self._converter = None

    def _get_converter(self):
        """Lazy-load the Docling converter (heavy import)."""
        if self._converter is None:
            from docling.document_converter import DocumentConverter
            self._converter = DocumentConverter()
        return self._converter

    async def process(self, file_path: str, filename: str) -> DoclingResult:
        """Process a document file and return structured chunks + tables.
        
        Runs Docling in a thread pool to avoid blocking the event loop.
        """
        return await asyncio.to_thread(self._process_sync, file_path, filename)

    def _process_sync(self, file_path: str, filename: str) -> DoclingResult:
        """Synchronous document processing with Docling."""
        try:
            converter = self._get_converter()
            result = converter.convert(file_path)
            doc = result.document

            chunks = []
            tables = []

            # Extract title
            title = None
            if hasattr(doc, 'title') and doc.title:
                title = str(doc.title)

            # ─── Export as Markdown and chunk by sections ──────────
            md_text = doc.export_to_markdown()

            if md_text and md_text.strip():
                chunks = self._chunk_markdown(md_text, filename)
            else:
                # Fallback: use plain text
                plain_text = doc.export_to_text() if hasattr(doc, 'export_to_text') else ""
                if plain_text and plain_text.strip():
                    chunks = self._chunk_plain_text(plain_text, filename)

            # ─── Extract tables ───────────────────────────────────
            if hasattr(doc, 'tables') and doc.tables:
                for table in doc.tables:
                    extracted = self._extract_table(table, filename)
                    if extracted:
                        tables.append(extracted)

            # Count pages
            num_pages = 0
            if hasattr(doc, 'pages'):
                num_pages = len(doc.pages) if doc.pages else 0

            logger.info(
                f"Docling processed '{filename}': {len(chunks)} chunks, "
                f"{len(tables)} tables, {num_pages} pages"
            )

            return DoclingResult(
                chunks=chunks,
                tables=tables,
                title=title,
                num_pages=num_pages,
            )

        except Exception as e:
            logger.error(f"Docling processing failed for '{filename}': {e}", exc_info=True)
            # Fallback to basic text extraction
            return self._fallback_extract(file_path, filename)

    def _chunk_markdown(self, md_text: str, filename: str) -> List[DocChunk]:
        """Chunk Markdown text by sections with overlap."""
        chunks = []

        # Split by headings (##, ###, etc.)
        sections = re.split(r'\n(?=#{1,4}\s)', md_text)

        current_section = None
        current_page = None
        buffer = []
        buffer_len = 0
        max_chunk_size = 1000
        overlap_size = 200

        for section in sections:
            # Extract heading if present
            heading_match = re.match(r'^(#{1,4})\s+(.*)', section)
            if heading_match:
                current_section = heading_match.group(2).strip()

            # Extract page number if present in metadata
            page_match = re.search(r'<!-- page[:\s]+(\d+)', section, re.IGNORECASE)
            if page_match:
                current_page = int(page_match.group(1))

            # Split long sections into smaller chunks
            paragraphs = section.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                if buffer_len + len(para) > max_chunk_size and buffer:
                    # Emit current buffer as a chunk
                    chunk_text = "\n\n".join(buffer)
                    chunks.append(DocChunk(
                        text=chunk_text,
                        source=filename,
                        page=current_page,
                        section=current_section,
                        chunk_type="text",
                    ))

                    # Keep overlap: last portion of text
                    overlap_text = chunk_text[-overlap_size:] if len(chunk_text) > overlap_size else ""
                    buffer = [overlap_text] if overlap_text else []
                    buffer_len = len(overlap_text)

                buffer.append(para)
                buffer_len += len(para)

        # Emit remaining buffer
        if buffer:
            chunk_text = "\n\n".join(buffer)
            if chunk_text.strip():
                chunks.append(DocChunk(
                    text=chunk_text,
                    source=filename,
                    page=current_page,
                    section=current_section,
                    chunk_type="text",
                ))

        # Number the chunks
        for i, chunk in enumerate(chunks):
            chunk.chunk_index = i
            chunk.total_chunks = len(chunks)

        return chunks

    def _chunk_plain_text(self, text: str, filename: str) -> List[DocChunk]:
        """Fallback: chunk plain text with simple splitting."""
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        split_texts = splitter.split_text(text)

        return [
            DocChunk(
                text=t,
                source=filename,
                chunk_type="text",
                chunk_index=i,
                total_chunks=len(split_texts),
            )
            for i, t in enumerate(split_texts)
        ]

    def _extract_table(self, table, filename: str) -> Optional[ExtractedTable]:
        """Extract a Docling table object into rows/columns."""
        try:
            # Docling tables have a .export_to_dataframe() method
            if hasattr(table, 'export_to_dataframe'):
                df = table.export_to_dataframe()
                if df is not None and not df.empty:
                    headers = [str(c) for c in df.columns]
                    rows = [[str(v) for v in row] for _, row in df.iterrows()]
                    caption = str(table.caption) if hasattr(table, 'caption') and table.caption else None

                    return ExtractedTable(
                        headers=headers,
                        rows=rows,
                        caption=caption,
                        source=filename,
                    )
            elif hasattr(table, 'to_dataframe'):
                df = table.to_dataframe()
                if df is not None and not df.empty:
                    headers = [str(c) for c in df.columns]
                    rows = [[str(v) for v in row] for _, row in df.iterrows()]
                    return ExtractedTable(
                        headers=headers,
                        rows=rows,
                        source=filename,
                    )
        except Exception as e:
            logger.warning(f"Failed to extract table from '{filename}': {e}")

        return None

    def _fallback_extract(self, file_path: str, filename: str) -> DoclingResult:
        """Fallback extraction when Docling fails — use basic libraries."""
        ext = os.path.splitext(filename)[1].lower()
        text = ""

        try:
            if ext == ".pdf":
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
            elif ext == ".docx":
                import docx
                doc = docx.Document(file_path)
                text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            elif ext in (".txt", ".md"):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            else:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
        except Exception as e:
            logger.error(f"Fallback extraction also failed for '{filename}': {e}")
            text = ""

        chunks = self._chunk_plain_text(text, filename) if text.strip() else []

        return DoclingResult(
            chunks=chunks,
            tables=[],
            title=None,
            num_pages=0,
        )


# ──────────────────────────────────────────────────────
# Singleton
# ──────────────────────────────────────────────────────

_processor: Optional[DoclingProcessor] = None


def get_docling_processor() -> DoclingProcessor:
    global _processor
    if _processor is None:
        _processor = DoclingProcessor()
    return _processor

