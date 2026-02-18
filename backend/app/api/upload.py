"""
Unified upload endpoint — smart auto-routing.

CSV/Excel → auto-detect columns → SQLite (tagged with collection)
PDF/DOCX/TXT/MD → Qdrant document embeddings (in the collection)

Handles both "tall" format (date column + value columns) and
"wide" format (label column + year/date columns as headers).

Saves a persistent copy of every upload to data/raw/uploads/.
"""

import json
import logging
import os
import re
import shutil
import tempfile
from datetime import datetime, timezone
from typing import Optional, List, Tuple

import pandas as pd
from fastapi import APIRouter, HTTPException, UploadFile, File, Form

from app.config import get_settings
from app.services.connectors.base import ParsedObservation
from app.services.ingestion import get_ingestion_service
from app.services.vectorstore import get_vectorstore_service
from app.core.database import get_session_factory
from app.models.timeseries import SeriesCatalog

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/upload", tags=["Upload"])

TIMESERIES_EXTENSIONS = {".csv", ".xlsx", ".xls"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
ALL_EXTENSIONS = TIMESERIES_EXTENSIONS | DOCUMENT_EXTENSIONS

# Internal Qdrant collections that should never receive user data
INTERNAL_COLLECTIONS = {"series_catalog_embeddings"}

# Where to persist uploaded files
RAW_UPLOAD_DIR = os.path.join(
    os.path.dirname(__file__), "..", "..", "data", "raw", "uploads"
)


# ──────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────

def _slugify(text: str) -> str:
    """Convert text to a clean slug for series_id."""
    s = text.strip().lower()
    s = re.sub(r"[^a-z0-9]+", "_", s)
    s = s.strip("_")
    return s


def _save_persistent_copy(content: bytes, filename: str, collection_name: str) -> str:
    """Save a persistent copy of the uploaded file under data/raw/uploads/{collection}/{timestamp}_{filename}."""
    dest_dir = os.path.join(RAW_UPLOAD_DIR, _slugify(collection_name))
    os.makedirs(dest_dir, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S")
    safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", filename)
    dest_path = os.path.join(dest_dir, f"{ts}_{safe_name}")
    with open(dest_path, "wb") as f:
        f.write(content)
    logger.info(f"Saved persistent copy: {dest_path}")
    return dest_path


def _looks_like_year(val) -> bool:
    """Check if a value looks like a year (1900-2100)."""
    try:
        n = int(float(val))
        return 1900 <= n <= 2100
    except (ValueError, TypeError):
        return False


def _column_headers_are_dates(columns: List) -> Tuple[bool, str]:
    """Check if column headers look like years or dates (wide-format indicator).
    
    Returns (is_wide, label_col) where label_col is the name column.
    """
    # Skip the first column (label column) and check the rest
    if len(columns) < 3:
        return False, ""

    # Check if columns[1:] look like years
    year_count = 0
    for col in columns[1:]:
        col_str = str(col).strip()
        if _looks_like_year(col_str):
            year_count += 1
        else:
            # Try parsing as date
            try:
                pd.to_datetime(col_str)
                year_count += 1
            except Exception:
                pass

    # If most of the non-label columns look like dates/years → wide format
    ratio = year_count / max(len(columns) - 1, 1)
    if ratio >= 0.5 and year_count >= 2:
        return True, str(columns[0])

    return False, ""


def _transpose_wide_to_tall(df: pd.DataFrame, label_col: str) -> pd.DataFrame:
    """Convert wide-format (label × years) to tall format (date, series, value).
    
    Wide:   | Sector     | 2005 | 2006 | 2007 |
            | Building   | 0.5  | 0.6  | 0.7  |
            | Transport  | 0.9  | 1.0  | 1.1  |
    
    Tall:   | date       | series_name | value |
            | 2005-01-01 | Building    | 0.5   |
            | 2005-01-01 | Transport   | 0.9   |
            | 2006-01-01 | Building    | 0.6   |
            ...
    """
    date_cols = [c for c in df.columns if c != label_col]

    # Melt wide → tall
    melted = df.melt(id_vars=[label_col], value_vars=date_cols,
                     var_name="date_raw", value_name="value")

    # Parse dates from column headers
    def parse_date(val):
        val_str = str(val).strip()
        if _looks_like_year(val_str):
            return pd.Timestamp(year=int(float(val_str)), month=1, day=1)
        try:
            return pd.to_datetime(val_str)
        except Exception:
            return pd.NaT

    melted["date"] = melted["date_raw"].apply(parse_date)
    melted = melted.dropna(subset=["date"])

    # Convert value to numeric
    melted["value"] = pd.to_numeric(melted["value"], errors="coerce")
    melted = melted.dropna(subset=["value"])

    melted["series_name"] = melted[label_col].astype(str).str.strip()
    melted = melted[["date", "series_name", "value"]]

    logger.info(f"Transposed wide→tall: {len(melted)} observations, "
                f"{melted['series_name'].nunique()} series")
    return melted


def _detect_date_column(df: pd.DataFrame) -> Optional[str]:
    """Auto-detect the most likely date column in a tall-format DataFrame."""
    # 1. Check column names that suggest dates
    date_keywords = ["date", "period", "month", "year", "time", "day", "timestamp"]
    for col in df.columns:
        if any(kw in str(col).lower() for kw in date_keywords):
            try:
                pd.to_datetime(df[col].dropna().head(10))
                return col
            except Exception:
                # Maybe it's integer years
                sample = df[col].dropna().head(10)
                if all(_looks_like_year(v) for v in sample):
                    return col
                continue

    # 2. Try the first column (common pattern)
    try:
        first_col = df.columns[0]
        sample = df[first_col].dropna().head(10)
        # Check for integer years first
        if all(_looks_like_year(v) for v in sample):
            return first_col
        pd.to_datetime(sample)
        return first_col
    except Exception:
        pass

    # 3. Try all columns
    for col in df.columns:
        try:
            sample = df[col].dropna().head(10)
            if all(_looks_like_year(v) for v in sample):
                return col
            pd.to_datetime(sample)
            return col
        except Exception:
            continue

    return None


def _parse_dates(series: pd.Series) -> pd.Series:
    """Parse dates flexibly — handles full dates, year-month, and plain years."""
    sample = series.dropna().head(10)

    # Check if values are integer years
    if all(_looks_like_year(v) for v in sample):
        return series.apply(
            lambda v: pd.Timestamp(year=int(float(v)), month=1, day=1)
            if pd.notna(v) and _looks_like_year(v) else pd.NaT
        )

    # Standard date parsing
    return pd.to_datetime(series, errors="coerce")


def _detect_numeric_columns(df: pd.DataFrame, date_col: str) -> list:
    """Find all numeric columns (excluding the date column)."""
    numeric_cols = []
    for col in df.columns:
        if col == date_col:
            continue
        try:
            numeric_data = pd.to_numeric(df[col], errors="coerce")
            if numeric_data.notna().sum() > 0:
                total_non_null = df[col].notna().sum()
                ratio = numeric_data.notna().sum() / total_non_null if total_non_null > 0 else 0
                if ratio > 0.5:
                    numeric_cols.append(col)
        except Exception:
            continue
    return numeric_cols


async def _auto_register_series(
    series_id: str,
    name: str,
    collection_name: str,
    filename: str,
):
    """Auto-create a series_catalog entry if it doesn't exist."""
    factory = await get_session_factory()
    session = factory()
    try:
        from sqlalchemy import select
        existing = await session.execute(
            select(SeriesCatalog).where(SeriesCatalog.series_id == series_id)
        )
        if existing.scalar():
            return  # Already exists

        new_series = SeriesCatalog(
            series_id=series_id,
            name=name,
            domain=collection_name,
            source="Upload",
            dataset=filename,
            description=f"Auto-imported from {filename}, column: {name}",
            synonyms=json.dumps([name.lower(), series_id.replace("_", " ")]),
            collection_name=collection_name,
        )
        session.add(new_series)
        await session.commit()
        logger.info(f"Auto-registered series: {series_id} in {collection_name}")
    except Exception as e:
        logger.warning(f"Failed to register series {series_id}: {e}")
        await session.rollback()
    finally:
        await session.close()


# ──────────────────────────────────────────────────────
# Main endpoint
# ──────────────────────────────────────────────────────

@router.post("/")
async def unified_upload(
    file: UploadFile = File(...),
    collection_name: str = Form(...),
    source: Optional[str] = Form(None),
):
    """
    Smart upload endpoint.

    - CSV/Excel → auto-detects structure (tall or wide) → stores in SQLite
    - PDF/DOCX/TXT/MD → stores document embeddings in Qdrant

    Saves a persistent copy of every upload to data/raw/uploads/.
    """
    # Block uploads to internal collections
    if collection_name in INTERNAL_COLLECTIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Cannot upload to internal collection '{collection_name}'. "
                   "Please create a new database first.",
        )

    ext = os.path.splitext(file.filename)[1].lower()

    if ext not in ALL_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(ALL_EXTENSIONS))}",
        )

    settings = get_settings()
    max_bytes = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
    content = await file.read()
    if len(content) > max_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"File too large ({len(content)/1024/1024:.1f} MB). Max: {settings.MAX_UPLOAD_SIZE_MB} MB"
        )

    # Save a persistent copy
    try:
        saved_path = _save_persistent_copy(content, file.filename, collection_name)
        logger.info(f"Persistent copy saved: {saved_path}")
    except Exception as e:
        logger.warning(f"Could not save persistent copy: {e}")

    try:
        # ─── DOCUMENT PATH ─────────────────────────────────
        if ext in DOCUMENT_EXTENSIONS:
            return await _handle_document_upload(content, file.filename, collection_name, ext)

        # ─── TIME SERIES PATH ──────────────────────────────
        return await _handle_timeseries_upload(content, file.filename, collection_name, source, ext)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


# ──────────────────────────────────────────────────────
# Document upload → Qdrant
# ──────────────────────────────────────────────────────

async def _handle_document_upload(
    content: bytes, filename: str, collection_name: str, ext: str
):
    """Process document → chunk → embed → store in Qdrant."""
    from langchain_core.documents import Document
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        # Read document text
        if ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(tmp_path)
                text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
            except ImportError:
                text = content.decode("utf-8", errors="ignore")
        elif ext in (".txt", ".md"):
            text = content.decode("utf-8", errors="ignore")
        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(tmp_path)
                text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError:
                text = content.decode("utf-8", errors="ignore")
        else:
            text = content.decode("utf-8", errors="ignore")

        if not text.strip():
            raise HTTPException(status_code=400, detail="No text content found in document")

        # Chunk the text
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        chunks = splitter.split_text(text)

        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": filename,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "collection": collection_name,
                },
            )
            for i, chunk in enumerate(chunks)
        ]

        # Store in Qdrant
        vs = get_vectorstore_service()
        num_added = await vs.add_documents(collection_name, documents)

        return {
            "type": "document",
            "filename": filename,
            "collection_name": collection_name,
            "chunks_created": num_added,
            "status": "success",
        }
    finally:
        os.unlink(tmp_path)


# ──────────────────────────────────────────────────────
# Time series upload → SQLite
# ──────────────────────────────────────────────────────

async def _handle_timeseries_upload(
    content: bytes, filename: str, collection_name: str,
    source: Optional[str], ext: str
):
    """Smart CSV/Excel processing — handles both tall and wide formats."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        # Read file
        if ext == ".csv":
            df = pd.read_csv(tmp_path)
        else:
            df = pd.read_excel(tmp_path)

        if df.empty:
            raise HTTPException(status_code=400, detail="File is empty")

        logger.info(f"Read {filename}: {df.shape[0]} rows × {df.shape[1]} cols. "
                    f"Columns: {list(df.columns)}")

        # ─── DETECT FORMAT: WIDE or TALL ───────────────
        is_wide, label_col = _column_headers_are_dates(df.columns)

        if is_wide:
            logger.info(f"Detected WIDE format (label col: '{label_col}')")
            return await _handle_wide_format(df, label_col, filename, collection_name, source)
        else:
            logger.info("Detected TALL format")
            return await _handle_tall_format(df, filename, collection_name, source)
    finally:
        os.unlink(tmp_path)


async def _handle_wide_format(
    df: pd.DataFrame, label_col: str,
    filename: str, collection_name: str, source: Optional[str]
):
    """Handle wide-format data (rows=series, columns=dates/years)."""
    melted = _transpose_wide_to_tall(df, label_col)

    if melted.empty:
        raise HTTPException(status_code=400, detail="No valid data after transposing wide-format file")

    observations = []
    series_names = melted["series_name"].unique()
    series_created = []

    for sname in series_names:
        raw_slug = _slugify(sname)
        series_id = f"{_slugify(collection_name)}__{raw_slug}"

        await _auto_register_series(
            series_id=series_id,
            name=sname,
            collection_name=collection_name,
            filename=filename,
        )
        series_created.append({"series_id": series_id, "column": sname})

        subset = melted[melted["series_name"] == sname]
        for _, row in subset.iterrows():
            try:
                observations.append(ParsedObservation(
                    series_id=series_id,
                    date=row["date"].date(),
                    value=float(row["value"]),
                ))
            except (ValueError, TypeError, AttributeError):
                continue

    if not observations:
        raise HTTPException(status_code=400, detail="No valid data rows found in wide-format file")

    # Load into SQLite
    service = get_ingestion_service()
    run_id = await service._create_run(source or "Upload", filename)
    load_result = await service.load_observations(observations, run_id)
    await service._finish_run(
        run_id=run_id,
        status="success",
        rows_inserted=load_result["inserted"],
        rows_updated=load_result["updated"],
        rows_skipped=load_result["skipped"],
        file_path=filename,
    )

    return {
        "type": "timeseries",
        "filename": filename,
        "collection_name": collection_name,
        "format_detected": "wide",
        "date_columns": [str(c) for c in df.columns if c != series_created[0]["column"]] if series_created else [],
        "series_detected": series_created,
        "rows_inserted": load_result["inserted"],
        "rows_updated": load_result["updated"],
        "rows_skipped": load_result["skipped"],
        "total": load_result["total"],
        "status": "success",
    }


async def _handle_tall_format(
    df: pd.DataFrame, filename: str,
    collection_name: str, source: Optional[str]
):
    """Handle tall-format data (standard: date col + value cols)."""
    # Auto-detect date column
    date_col = _detect_date_column(df)
    if not date_col:
        raise HTTPException(
            status_code=400,
            detail=f"Could not detect a date column. Found columns: {list(df.columns)}. "
                   "Please ensure one column contains dates (or use years as column headers for wide format)."
        )

    # Auto-detect numeric columns
    numeric_cols = _detect_numeric_columns(df, date_col)
    if not numeric_cols:
        raise HTTPException(
            status_code=400,
            detail=f"No numeric columns found (besides date column '{date_col}'). "
                   f"Found columns: {list(df.columns)}"
        )

    # Parse dates (handles years, year-month, full dates)
    df[date_col] = _parse_dates(df[date_col])
    df = df.dropna(subset=[date_col])

    # Build observations and auto-register series
    observations = []
    series_created = []

    for col in numeric_cols:
        raw_slug = _slugify(col)
        series_id = f"{_slugify(collection_name)}__{raw_slug}"

        await _auto_register_series(
            series_id=series_id,
            name=col.strip(),
            collection_name=collection_name,
            filename=filename,
        )
        series_created.append({"series_id": series_id, "column": col})

        # Convert column to numeric
        numeric_data = pd.to_numeric(df[col], errors="coerce")

        for _, row in df.iterrows():
            try:
                date_val = row[date_col].date()
                value = float(numeric_data[row.name]) if pd.notna(numeric_data[row.name]) else None
                if value is not None:
                    observations.append(ParsedObservation(
                        series_id=series_id,
                        date=date_val,
                        value=value,
                    ))
            except (ValueError, TypeError, AttributeError):
                continue

    if not observations:
        raise HTTPException(status_code=400, detail="No valid data rows found in file")

    # Load into SQLite
    service = get_ingestion_service()
    run_id = await service._create_run(source or "Upload", filename)
    load_result = await service.load_observations(observations, run_id)
    await service._finish_run(
        run_id=run_id,
        status="success",
        rows_inserted=load_result["inserted"],
        rows_updated=load_result["updated"],
        rows_skipped=load_result["skipped"],
        file_path=filename,
    )

    return {
        "type": "timeseries",
        "filename": filename,
        "collection_name": collection_name,
        "format_detected": "tall",
        "date_column": date_col,
        "series_detected": series_created,
        "rows_inserted": load_result["inserted"],
        "rows_updated": load_result["updated"],
        "rows_skipped": load_result["skipped"],
        "total": load_result["total"],
        "status": "success",
    }
